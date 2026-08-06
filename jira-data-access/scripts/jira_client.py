#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
JIRA数据访问客户端
功能：从JIRA系统提取缺陷数据，支持查询、过滤、导出等操作
"""

import urllib.parse
import json
import time
import re
import urllib.request
import os
import sys
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from bs4 import BeautifulSoup
import datetime
import requests
from urllib.parse import unquote, quote, urlparse, parse_qs, urlunparse

# docx相关导入
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_SUPPORTED = True
except ImportError:
    DOCX_SUPPORTED = False
    print("⚠️  DOCX功能未启用，请安装python-docx库：pip install python-docx")

def _has_status_clause(jql):
    """粗略检测 JQL 中是否已经包含 status 相关条件（避免重复注入）。"""
    if not jql:
        return False
    # 先去掉 ORDER BY 子句，避免 ORDER BY status 被误识别为 status 条件
    stripped = re.sub(r'\bORDER\s+BY\b.*', '', jql, flags=re.IGNORECASE)
    # 再去掉引号内的文本，避免 text ~ "status" 这类误命中
    stripped = re.sub(r'"[^"]*"', '""', stripped)
    stripped = re.sub(r"'[^']*'", "''", stripped)
    return bool(re.search(r'\bstatus\b', stripped, re.IGNORECASE))


def normalize_jql(jql, exclude_closed=True):
    """
    规范化 JQL：默认自动追加 `AND status != "ST_Closed"`（当 JQL 中未显式写 status 时）。
    这样「一般搜索问题，只搜索非 ST_Closed 的问题」；用户要搜关闭的可显式在 JQL 里写 status，
    或调用时传 exclude_closed=False。
    会自动处理 ORDER BY，保证注入的过滤条件位于 ORDER BY 之前。
    :return: (normalized_jql, closed_excluded)
    """
    has_status = _has_status_clause(jql)
    if not exclude_closed or has_status:
        return jql, has_status
    # 如果有 ORDER BY，把过滤条件插入到 ORDER BY 之前
    order_match = re.search(r'\bORDER\s+BY\b', jql, re.IGNORECASE)
    if order_match:
        before = jql[:order_match.start()].strip()
        after = jql[order_match.start():].strip()
        # 去掉 before 末尾可能残留的 AND/OR（防御性处理）
        before = re.sub(r'\s+(AND|OR)\s*$', '', before, flags=re.IGNORECASE)
        if before:
            return f"({before}) AND status != \"ST_Closed\" {after}", True
        return f"{after} AND status != \"ST_Closed\"", True  # 仅 ORDER BY 的极端情况
    return f"({jql}) AND status != \"ST_Closed\"", True


class JIRAClient:
    """
    JIRA客户端，负责与JIRA系统交互，提取和处理数据
    """
    
    def __init__(self, account, password, base_url="https://ticket.ikotek.com"):
        """
        初始化JIRA客户端
        :param account: JIRA账号
        :param password: JIRA密码
        :param base_url: JIRA系统基础URL
        """
        self.account = account
        self.password = password
        self.base_url = base_url.rstrip('/')
        self.session = requests.session()
        self.is_logged_in = False
        
    def login(self):
        """登录JIRA系统"""
        url = f"{self.base_url}/login.jsp"
        payload = f'os_username={urllib.parse.quote(self.account)}&os_password={urllib.parse.quote(self.password)}&os_destination=https%3A%2F%2Fticket.ikotek.com%2Fissues%2F%3Fjql%3D&user_role=&atl_token=&login=Log%2BIn'
        headers = {
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.7',
            'Content-Type': 'application/x-www-form-urlencoded',
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/127.0.0.0 Safari/537.36 Edg/127.0.0.0'
        }
        response = self.session.post(url, headers=headers, data=payload)
        self.is_logged_in = response.status_code == 200
        return self.is_logged_in
    
    def _remove_url_parameter(self, url, param_to_remove):
        """移除URL中的指定参数"""
        parsed_url = urlparse(url)
        query_params = parse_qs(parsed_url.query)
        if param_to_remove in query_params:
            del query_params[param_to_remove]
        new_query = '&'.join([f"{k}={v[0]}" for k, v in query_params.items()])
        new_url = urlunparse(parsed_url._replace(query=new_query))
        return new_url
    
    def query_by_jql(self, jql_query, exclude_closed=True):
        """
        通过JQL查询JIRA数据
        :param jql_query: JQL查询语句
        :param exclude_closed: 是否默认追加 status != "ST_Closed"（当 JQL 未显式写 status 时）
        :return: (headers, data) 表头列表和数据字典列表
        """
        if not self.is_logged_in:
            if not self.login():
                raise Exception("JIRA登录失败")
        
        jql_query, _ = normalize_jql(jql_query, exclude_closed=exclude_closed)
        encoded_jql = urllib.parse.quote(jql_query)
        url = f"{self.base_url}/sr/jira.issueviews:searchrequest-html-all-fields/temp/SearchRequest.html?jqlQuery={encoded_jql}"
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/127.0.0.0 Safari/537.36 Edg/127.0.0.0'
        }
        
        response = self.session.get(url, headers=headers)
        if response.status_code != 200:
            raise Exception(f"获取JIRA数据失败，状态码: {response.status_code}")
        
        return self._parse_html(response.content)
    
    def query_by_url(self, jira_url, exclude_closed=True):
        """
        通过JIRA搜索URL查询数据
        :param jira_url: JIRA搜索页面的完整URL
        :param exclude_closed: 是否默认追加 status != "ST_Closed"（当 JQL 未显式写 status 时）
        :return: (headers, data) 表头列表和数据字典列表
        """
        if not self.is_logged_in:
            if not self.login():
                raise Exception("JIRA登录失败")
        
        # 处理URL，移除filter参数
        cleaned_url = self._remove_url_parameter(jira_url, 'filter')
        
        # 提取JQL部分
        if 'jql=' in cleaned_url:
            jql_part = cleaned_url.split("jql=")[-1]
            decoded_jql = urllib.parse.unquote(jql_part)
            return self.query_by_jql(decoded_jql, exclude_closed=exclude_closed)
        else:
            raise Exception("URL中未找到JQL查询参数")
    
    def _parse_html(self, html_content):
        """解析JIRA HTML响应内容"""
        soup = BeautifulSoup(html_content, "html.parser")
        tr_list = soup.find_all("tr", {"class": "issuerow"})
        
        raw_records = []
        
        for tr in tr_list:
            try:
                record = {
                    "Bug号": tr.find("td", {"class": "issuekey"}).find("a")["data-issue-key"].strip(),
                    "链接": self.base_url + tr.find("td", {"class": "issuekey"}).find("a")["href"],
                    "Bug描述": tr.find("td", {"class": "summary"}).text.strip(),
                    "Bug状态": tr.find("td", {"class": "status"}).text.strip(),
                    "Bug严重等级": tr.find("td", {"class": "customfield_10203"}).text.strip(),
                    "Bug发现项目": tr.find("td", {"class": "customfield_10240"}).text.strip(),
                    "Bug发现版本": tr.find("td", {"class": "customfield_10249"}).text.strip() + "_" + tr.find("td", {"class": "customfield_10223"}).text.strip(),
                    "ST评审意见": tr.find("td", {"class": "customfield_10226"}).text.strip(),
                    "SW评估意见": tr.find("td", {"class": "customfield_10239"}).text.strip()
                }
                raw_records.append(record)
            except Exception as e:
                print(f"⚠️  解析JIRA数据时出错: {e}")
                continue
        
        # 设置原始表头
        headers = []
        if raw_records:
            headers = list(raw_records[0].keys())
        
        return headers, raw_records

    # ------------------------------------------------------------------
    # 单个JIRA问题详情访问（访问具体JIRA页面）
    # 接口数据来源于浏览器访问具体JIRA问题页面时抓取的网络请求
    # （见 references/jira_issue_page_interface.md）。
    # 采用 JIRA 标准 REST 接口 /rest/api/2/issue/{issueKey} 获取结构化数据。
    # ------------------------------------------------------------------
    def _ensure_field_names(self):
        """
        获取字段ID -> 中文名称映射（用于自定义字段展示），结果缓存到 self._field_names
        优先从 /rest/api/2/field 获取，失败时回退到内置映射
        """
        if getattr(self, '_field_names', None) is not None:
            return self._field_names
        self._field_names = {}
        try:
            if not self.is_logged_in:
                self.login()
            url = f"{self.base_url}/rest/api/2/field"
            resp = self.session.get(url, headers={
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/127.0.0.0 Safari/537.36',
                'Accept': 'application/json'
            })
            if resp.status_code == 200:
                for f in resp.json():
                    self._field_names[f['id']] = f.get('name', f['id'])
        except Exception as e:
            print(f"⚠️  获取字段名称映射失败，使用内置映射: {e}")
        # 内置兜底映射（来自已知JIRA配置及HAR抓包）
        fallback = {
            "customfield_10203": "BUG严重等级",
            "customfield_10226": "ST BUG评估意见",
            "customfield_10239": "SW BUG评估意见",
            "customfield_10240": "BUG发现的项目",
            "customfield_10249": "BUG发现的软件版本",
            "customfield_10223": "BUG发现的V版本",
            "customfield_10227": "BUG来源",
            "customfield_10221": "BUG优先级",
            "customfield_10246": "BUG关闭版本",
            "customfield_10250": "BUG关闭V版本",
        }
        for k, v in fallback.items():
            self._field_names.setdefault(k, v)
        return self._field_names

    def get_issue_detail(self, issue_key, fields=None):
        """
        访问并获取单个JIRA问题的完整详情（即具体JIRA问题页面数据）
        使用JIRA标准REST接口 /rest/api/2/issue/{issueKey}
        :param issue_key: 问题Key，如 'QDM565EA-396'
        :param fields: 需要返回的字段列表（可选，默认返回全部字段）
        :return: (raw_json, parsed_detail) 原始JSON与解析后的结构化字典
        """
        if not self.is_logged_in:
            if not self.login():
                raise Exception("JIRA登录失败")
        
        url = f"{self.base_url}/rest/api/2/issue/{issue_key}"
        params = {}
        if fields:
            params['fields'] = ','.join(fields)
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/127.0.0.0 Safari/537.36',
            'Accept': 'application/json'
        }
        
        response = self.session.get(url, params=params, headers=headers)
        if response.status_code != 200:
            raise Exception(f"获取JIRA问题详情失败，状态码: {response.status_code} - {response.text[:200]}")
        
        raw = response.json()
        return raw, self._parse_issue_detail(raw)

    # ---------------------- 取值辅助方法 ----------------------
    @staticmethod
    def _name_of(val):
        """从状态/优先级/类型等对象中提取名称"""
        if isinstance(val, dict):
            return val.get('name') or val.get('value') or ''
        return val or ''

    @staticmethod
    def _user_name(val):
        """从用户对象中提取显示名"""
        if isinstance(val, dict):
            return val.get('displayName') or val.get('name') or ''
        return val or ''

    @staticmethod
    def _adf_to_text(node):
        """将Atlassian Document Format(ADF)富文本节点转为纯文本"""
        if node is None:
            return ''
        if isinstance(node, str):
            return node
        if isinstance(node, dict):
            if node.get('type') == 'text':
                return node.get('text', '')
            parts = []
            for child in node.get('content', []):
                parts.append(JIRAClient._adf_to_text(child))
            return '\n'.join(p for p in parts if p)
        if isinstance(node, list):
            return '\n'.join(JIRAClient._adf_to_text(c) for c in node)
        return str(node)

    @staticmethod
    def _flatten_value(val):
        """将REST返回的字段值统一转换为可读字符串"""
        if val is None:
            return ''
        if isinstance(val, str):
            return val
        if isinstance(val, list):
            return '; '.join(JIRAClient._flatten_value(x) for x in val if x is not None)
        if isinstance(val, dict):
            if 'value' in val:            # 单选/多选自定义字段选项
                return val['value']
            if 'displayName' in val:      # 用户字段
                return val['displayName']
            if 'name' in val:             # 状态/优先级/类型等
                return val['name']
            if 'key' in val and 'summary' in val:  # 关联问题
                return f"{val['key']} {val.get('summary', '')}"
            if val.get('type') == 'doc':  # ADF富文本
                return JIRAClient._adf_to_text(val)
            return json.dumps(val, ensure_ascii=False)
        return str(val)

    def _parse_issue_detail(self, raw):
        """将单个问题的REST JSON解析为结构化字典"""
        names = self._ensure_field_names()
        f = raw.get('fields', {})
        
        detail = {
            'key': raw.get('key', ''),
            'summary': f.get('summary', ''),
            'issuetype': self._name_of(f.get('issuetype')),
            'status': self._name_of(f.get('status')),
            'priority': self._name_of(f.get('priority')),
            'reporter': self._user_name(f.get('reporter')),
            'assignee': self._user_name(f.get('assignee')),
            'created': f.get('created', ''),
            'updated': f.get('updated', ''),
            'description': self._flatten_value(f.get('description')),
            'components': [c.get('name', '') for c in f.get('components', [])],
            'fixVersions': [v.get('name', '') for v in f.get('fixVersions', [])],
            'labels': f.get('labels', []),
            'custom_fields': {},
            'comments': [],
            'attachments': [],
            'issuelinks': [],
        }
        
        # 自定义字段（customfield_ID -> 中文名称 -> 值）
        for cid, val in f.items():
            if cid.startswith('customfield_') and val is not None:
                label = names.get(cid, cid)
                detail['custom_fields'][label] = self._flatten_value(val)
        
        # 过滤系统噪声字段（全局公告栏、开发状态Java对象dump、排名序号等）
        NOISE_LABELS = {"提示信息", "请注意", "请注意！", "注意事项", "Development", "Request participants"}
        DEV_DUMP_MARK = "com.atlassian.jira.plugin.devstatus"
        import re as _re
        cleaned = {}
        for label, val in detail['custom_fields'].items():
            if label in NOISE_LABELS:
                continue
            if DEV_DUMP_MARK in str(val):
                continue
            # 排名序号字段（值形如 "0|i0dgrz:"）属于系统内部排序，对用户无意义
            if label == '等级' and _re.match(r'^\d+\|', str(val).strip()):
                continue
            cleaned[label] = val
        detail['custom_fields'] = cleaned
        
        # 评论
        for c in f.get('comment', {}).get('comments', []):
            detail['comments'].append({
                'author': self._user_name(c.get('author')),
                'created': c.get('created', ''),
                'body': self._flatten_value(c.get('body')),
            })
        
        # 附件
        for a in f.get('attachment', []):
            detail['attachments'].append({
                'filename': a.get('filename', ''),
                'created': a.get('created', ''),
                'url': a.get('content', ''),
            })
        
        # 关联问题
        for l in f.get('issuelinks', []):
            outward = l.get('outwardIssue')
            inward = l.get('inwardIssue')
            linked = outward or inward
            if linked:
                detail['issuelinks'].append({
                    'type': l.get('type', {}).get('name', ''),
                    'direction': 'outward' if outward else 'inward',
                    'key': linked.get('key', ''),
                    'summary': linked.get('summary', ''),
                })
        
        return detail

    # ==================================================================
    # 创建 JIRA 问题（Create Issue）
    # 接口来源：ticket.ikotek.com.txt（HAR 抓包）。流程：
    #   1) GET /secure/CreateIssue.jspa?pid={projectId}
    #        -> 从页面隐藏域提取 formToken 与 atl_token（与 cookie 可能不同，必须用表单里的）
    #   2) （可选）逐一对附件 POST /rest/internal/2/AttachTemporaryFile
    #        -> 原始文件字节作为 body，返回 {"id":"tempXXXX", ...}，即 filetoconvert 值
    #   3) POST /secure/QuickCreateIssue.jspa?decorator=none （form 表单编码）
    #        -> 返回 JSON，含 "issueKey"
    # 选项型字段（如 问题归属 / BUG严重等级）的值需要用「选项 ID」（数字），
    # 这些 ID 通过 /rest/api/2/issue/createmeta 解析（label -> id）。
    # ==================================================================
    def get_create_meta(self, project_key=None, issuetype_name=None, project_id=None):
        """
        获取创建元数据：项目 ID、问题类型 ID 与所有字段的 schema（含选项 ID）。
        :return: {'project_id':str, 'issuetype_id':str, 'fields':{fid: meta}}
        """
        if not self.is_logged_in:
            if not self.login():
                raise Exception("JIRA登录失败")
        url = f"{self.base_url}/rest/api/2/issue/createmeta"
        params = {"expand": "projects.issuetypes.fields"}
        if project_key:
            params["projectKeys"] = project_key
        if issuetype_name:
            params["issuetypeNames"] = issuetype_name
        if project_id:
            params["projectIds"] = str(project_id)
        resp = self.session.get(url, params=params, headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36',
            'Accept': 'application/json'
        })
        if resp.status_code != 200:
            raise Exception(f"获取创建元数据失败，状态码: {resp.status_code}")
        data = resp.json()
        proj = None
        for p in data.get("projects", []):
            if project_key and p.get("key") != project_key:
                if project_id and str(p.get("id")) != str(project_id):
                    continue
            proj = p
            break
        if proj is None and data.get("projects"):
            proj = data["projects"][0]
        if proj is None:
            raise Exception("createmeta 未返回任何项目，请检查 project_key/project_id")
        it = None
        for t in proj.get("issuetypes", []):
            if issuetype_name and t.get("name") != issuetype_name:
                continue
            it = t
            break
        if it is None and proj.get("issuetypes"):
            it = proj["issuetypes"][0]
        return {
            "project_id": proj.get("id"),
            "issuetype_id": it.get("id") if it else None,
            "fields": it.get("fields", {}) if it else {},
        }

    def _get_create_form_tokens(self, project_id):
        """GET 创建表单页，提取 formToken 与 atl_token（隐藏域，可能与 cookie 不同）"""
        url = f"{self.base_url}/secure/CreateIssue.jspa?pid={project_id}"
        resp = self.session.get(url, headers={
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'})
        if resp.status_code != 200:
            raise Exception(f"获取创建表单失败，状态码: {resp.status_code}")
        html = resp.text
        m = re.search(r'name="formToken"[^>]*?value="([^"]*)"', html, re.S)
        a = re.search(r'name="atl_token"[^>]*?value="([^"]*)"', html, re.S)
        return (m.group(1) if m else "", a.group(1) if a else "")

    @staticmethod
    def _find_option_id(allowed, value):
        """在 allowedValues 中按 id / value / name 解析出选项 ID（字符串）"""
        if value is None:
            return None
        if isinstance(value, int) or (isinstance(value, str) and value.isdigit()):
            return str(value)
        v = str(value)
        for o in allowed:
            if o.get("value") == v or o.get("name") == v or str(o.get("id")) == v:
                return str(o.get("id"))
        for o in allowed:
            if v.lower() in str(o.get("value", "")).lower():
                return str(o.get("id"))
        return None

    def _resolve_field_value(self, fid, meta, value):
        """
        将用户传入的字段值解析为 POST 表单的 (key, value) 列表。
        - 选项型：接受选项 ID（数字）或选项 label（中文），label 经 allowedValues 解析为 ID
        - 级联选择（cascading，如 功能类别）：接受 [parent, child]（可为 label 或 ID）
        - 多选 / 数组：接受 list，逐项解析
        - 文本 / 日期 / 数字：原样转为字符串
        """
        schema = meta.get("schema", {}) or {}
        ftype = schema.get("type")
        allowed = meta.get("allowedValues")
        is_cascading = (ftype in ("cascading", "option-with-child")
                        or (bool(allowed) and isinstance(allowed[0], dict) and bool(allowed[0].get("children"))))
        if isinstance(value, list) and not allowed and ftype != "cascading":
            return [(fid, str(v)) for v in value]
        if allowed:
            if is_cascading:
                pair = (list(value) + [None, None])[:2] if isinstance(value, list) else [value, None]
                parent, child = pair
                pid_ = self._find_option_id(allowed, parent)
                res = [(fid, pid_)] if pid_ else []
                if child and pid_:
                    for o in allowed:
                        if str(o.get("id")) == str(pid_):
                            children = o.get("children") or o.get("options") or []
                            cid = self._find_option_id(children, child)
                            if cid:
                                res.append((fid + ":1", cid))
                            break
                return res
            # 多选
            if isinstance(value, list):
                out = []
                for v in value:
                    x = self._find_option_id(allowed, v)
                    if x:
                        out.append((fid, x))
                if not out:
                    raise ValueError(f"字段「{meta.get('name')}」无法解析任何选项: {value!r}")
                return out
            oid = self._find_option_id(allowed, value)
            if oid is None:
                # 同时显示value和id，方便用户选择
                options_display = []
                for o in allowed[:20]:
                    val = o.get('value', '')
                    oid = o.get('id', '')
                    if val and oid:
                        options_display.append(f"{val} (ID: {oid})")
                    elif val:
                        options_display.append(val)
                raise ValueError(f"字段「{meta.get('name')}」无法解析选项: {value!r}\n可选值（前20个）：{options_display}")
            return [(fid, oid)]
        if isinstance(value, list):
            return [(fid, str(v)) for v in value]
        return [(fid, str(value))]

    def upload_attachment_temp(self, file_path, project_id, form_token, atl_token):
        """
        上传附件到临时区，返回 temp id（如 'temp123456'），供创建时 filetoconvert 使用。
        注意：该文件仅在随后成功创建问题时才会挂接到问题；未创建则为孤立临时文件。
        """
        with open(file_path, "rb") as f:
            content = f.read()
        fn = os.path.basename(file_path)
        url = f"{self.base_url}/rest/internal/2/AttachTemporaryFile"
        params = {
            "filename": fn,
            "size": str(len(content)),
            "atl_token": atl_token,
            "formToken": form_token,
            "projectId": str(project_id),
        }
        headers = {
            "Content-Type": "text/plain",
            "X-Requested-With": "XMLHttpRequest",
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
            "Origin": self.base_url,
        }
        resp = self.session.post(url, params=params, data=content, headers=headers, timeout=120)
        if resp.status_code not in (200, 201):
            raise Exception(f"附件上传失败，状态码: {resp.status_code} - {resp.text[:200]}")
        try:
            return resp.json()["id"]
        except Exception:
            raise Exception(f"附件上传响应解析失败: {resp.text[:200]}")

    def _prepare_create(self, project_key=None, issuetype_name="ST-BUG", summary=None,
                         description=None, fields=None, assignee=None, priority=None,
                         components=None, links=None, project_id=None):
        """
        只读地解析并组装创建 JIRA 所需的全部 POST 表单数据（不发起创建、不上传附件）。
        返回 dict: {pid, form_token, atl_token, post_items, name2id, meta}
        """
        if not summary:
            raise ValueError("summary（标题）为必填项")
        if not self.is_logged_in:
            if not self.login():
                raise Exception("JIRA登录失败")
        meta = self.get_create_meta(project_key=project_key, issuetype_name=issuetype_name, project_id=project_id)
        pid = meta["project_id"]
        iid = meta["issuetype_id"]
        if not pid or not iid:
            raise Exception("无法确定项目或问题类型（请检查 project_key / project_id / issuetype_name）")
        form_token, atl_token = self._get_create_form_tokens(pid)
        if not form_token:
            raise Exception("无法从创建表单获取 formToken")
        atl = atl_token or self.session.cookies.get("atlassian.xsrf.token", "")

        fields_meta = meta["fields"]
        name2id = {}
        for fid, m in fields_meta.items():
            if fid.startswith("customfield_") and m.get("name"):
                name2id[m["name"]] = fid

        post_items = [
            ("pid", str(pid)),
            ("issuetype", str(iid)),
            ("atl_token", atl),
            ("formToken", form_token),
            ("summary", summary),
            ("isCreateIssue", "true"),
            ("dnd-dropzone", ""),
            ("timetracking_originalestimate", ""),
            ("timetracking_remainingestimate", ""),
            ("hasWorkStarted", ""),
            ("duedate", ""),
        ]
        retain = set(["project", "issuetype", "summary", "priority", "components",
                      "assignee", "fixVersions", "duedate", "labels", "timetracking",
                      "issuelinks", "attachment"])

        # 描述 -> customfield_10244（该项目的“BUG描述”字段）
        desc_field = "customfield_10244"
        if description is not None:
            post_items.append((desc_field, description))
            retain.add(desc_field)

        def add_standard(std_name, value):
            if value is None:
                return
            m = fields_meta.get(std_name)
            if not m:
                post_items.append((std_name, str(value)))
                retain.add(std_name)
                return
            for k, v in self._resolve_field_value(std_name, m, value):
                post_items.append((k, v))
            retain.add(std_name)

        # QDM551项目优先级默认映射：Major -> ID 10003，避免名称解析失败
        if priority and project_key == "IKQDM551EA" and priority.lower() == "major":
            add_standard("priority", "10003")
        else:
            add_standard("priority", priority)
        add_standard("components", components)
        if assignee:
            post_items.append(("assignee", assignee))
            retain.add("assignee")

        if fields:
            for key, val in fields.items():
                if val is None:
                    continue
                fid = key if str(key).startswith("customfield_") else name2id.get(key)
                if not fid or fid not in fields_meta:
                    print(f"WARNING 未知字段，已忽略: {key}")
                    continue
                if fid == desc_field and description is not None:
                    continue  # 已被 description 处理
                for k, v in self._resolve_field_value(fid, fields_meta[fid], val):
                    post_items.append((k, v))
                retain.add(fid)

        if links:
            post_items.append(("issuelinks", "issuelinks"))
            retain.add("issuelinks")
            for i, ln in enumerate(links):
                if isinstance(ln, dict):
                    ltype = ln.get("type", "blocks")
                    lkey = ln.get("key")
                else:
                    ltype, lkey = "blocks", ln
                post_items.append((f"issuelinks-linktype", ltype))
                if lkey:
                    post_items.append((f"issuelinks-issueLink-{i}", lkey))

        for r in retain:
            post_items.append(("fieldsToRetain", r))

        return {
            "pid": pid, "form_token": form_token, "atl_token": atl,
            "post_items": post_items, "name2id": name2id, "meta": meta,
        }

    def build_create_draft(self, project_key=None, issuetype_name="ST-BUG", summary=None,
                           description=None, fields=None, assignee=None, priority=None,
                           components=None, attachments=None, links=None, project_id=None):
        """
        只读地生成创建草稿，不发起任何写操作（不创建、不上传附件）。
        用于「先确认再创建」流程：把将要提交的字段全部解析好并展示给用户确认。

        :return: dict{
            summary, description,
            resolved_fields: {字段名: 解析后值（人类可读，已解析为选项 ID/label）},
            post_items_preview: [(表单字段名, 值)...] 实际将提交的表单数据,
            attachments: [本地路径...]（仅列出，未上传）,
            links,
        }
        """
        prep = self._prepare_create(
            project_key=project_key, issuetype_name=issuetype_name, summary=summary,
            description=description, fields=fields, assignee=assignee, priority=priority,
            components=components, links=links, project_id=project_id)

        # 人类可读的字段摘要（过滤掉系统内部字段）
        _skip = {"pid", "issuetype", "atl_token", "formToken", "isCreateIssue",
                 "dnd-dropzone", "timetracking_originalestimate",
                 "timetracking_remainingestimate", "hasWorkStarted", "duedate"}
        readable = {}
        for k, v in prep["post_items"]:
            if k in _skip or k.startswith("fieldsToRetain"):
                continue
            if k in readable:
                readable[k] = readable[k] if isinstance(readable[k], list) else [readable[k]]
                readable[k].append(v)
            else:
                readable[k] = v

        return {
            "summary": summary,
            "description": description,
            "resolved_fields": readable,
            "post_items_preview": prep["post_items"],
            "attachments": attachments or [],
            "links": links or [],
        }

    def create_issue(self, project_key=None, issuetype_name="ST-BUG", summary=None,
                     description=None, fields=None, assignee=None, priority=None,
                     components=None, attachments=None, links=None, project_id=None,
                     draft=False, auto_complete=True, verify_create=True):
        """
        创建 JIRA 问题，并返回新问题的 Key（如 'IKQDM551EA-700'）。

        :param project_key: 项目 Key，如 'IKQDM551EA'（与 project_id 二选一）
        :param issuetype_name: 问题类型名称，默认 'ST-BUG'
        :param summary: 标题（必填）
        :param description: 描述（映射到自定义字段 customfield_10244 = BUG描述）
        :param fields: 其它自定义字段字典。key 可为 customfield_XXXXX 或中文名称；
                       value 对选项型字段可为「选项 ID」或「选项中文 label」，
                       对级联选择（功能类别）为 [parent, child]，多选为 list。
        :param assignee: 处理人邮箱，如 'siliver.nong@ikotek.com'
        :param priority: 优先级（label 或 ID，如 'Major'/'10003'）
        :param components: 模块/组件（label 或 ID，如 'FM33FK545_APP'/10700）
        :param attachments: 本地附件文件路径列表
        :param links: 关联问题列表，每项 {'type':'blocks','key':'IKQDM551EA-600'} 或 ('blocks','IKQDM551EA-600')
        :param draft: 若为 True，仅生成草稿（只读、不创建、不上传附件），
                      返回 build_create_draft 的结构；用于「先确认再创建」流程。默认 False。
        :param auto_complete: 自动补全缺失的必填字段，默认 True。开启后会自动查询最近3个月同类工单的默认值
        :param verify_create: 创建后验证工单是否真实存在，默认 True。开启后会调用get_issue_detail验证
        :return: draft=True 时返回草稿 dict；否则返回新问题 Key 字符串
        """
        if draft:
            return self.build_create_draft(
                project_key=project_key, issuetype_name=issuetype_name, summary=summary,
                description=description, fields=fields, assignee=assignee, priority=priority,
                components=components, attachments=attachments, links=links, project_id=project_id)

        # 自动补全缺失的必填字段
        if auto_complete and not draft:
            try:
                # 查询最近3个月的同类工单获取默认值
                jql = f"project = {project_key} AND issuetype = '{issuetype_name}' AND created >= -90d ORDER BY created DESC"
                issues = self.search_issues(jql, max_results=10)
                if issues:
                    # 取最新的工单作为参考
                    latest_issue = issues[0]
                    _, latest_detail = self.get_issue_detail(latest_issue['key'])
                    latest_fields = latest_detail['custom_fields']
                    
                    # 定义需要自动补全的必填字段列表
                    required_fields = ['ST BUG评估意见', '项目阶段', 'BUG发现的项目', 
                                     'BUG发现的软件版本', 'BUG优先级', '功能类别', 
                                     'BUG发现的A版本', 'BUG发现的V版本']
                    
                    fields = fields or {}
                    for field in required_fields:
                        # 如果用户未提供该字段，且历史工单有值，则自动补全
                        if field not in fields and field in latest_fields:
                            fields[field] = latest_fields[field]
                            print(f"[自动补全] 字段「{field}」已设置为历史默认值: {fields[field]}")
            except Exception as e:
                print(f"[自动补全警告] 无法获取历史工单默认值: {str(e)}")

        prep = self._prepare_create(
            project_key=project_key, issuetype_name=issuetype_name, summary=summary,
            description=description, fields=fields, assignee=assignee, priority=priority,
            components=components, links=links, project_id=project_id)
        post_items = prep["post_items"]
        pid = prep["pid"]; form_token = prep["form_token"]; atl = prep["atl_token"]

        temp_ids = []
        if attachments:
            for fp in attachments:
                tid = self.upload_attachment_temp(fp, pid, form_token, atl)
                temp_ids.append(tid)
                post_items.append(("filetoconvert", tid))

        url = f"{self.base_url}/secure/QuickCreateIssue.jspa?decorator=none"
        headers = {
            "Content-Type": "application/x-www-form-urlencoded; charset=UTF-8",
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/127.0.0.0 Safari/537.36 Edg/127.0.0.0",
            "X-Requested-With": "XMLHttpRequest",
            "Origin": self.base_url,
            "Referer": f"{self.base_url}/issues/?jql=",
            "Accept": "*/*",
        }
        resp = self.session.post(url, data=post_items, headers=headers)
        if resp.status_code not in (200, 201):
            raise Exception(f"创建JIRA失败，状态码: {resp.status_code} - {resp.text[:500]}")
        try:
            j = resp.json()
        except Exception:
            raise Exception(f"创建JIRA响应解析失败: {resp.text[:500]}")
        key = j.get("issueKey") or j.get("key")
        if not key:
            raise Exception(f"创建JIRA未返回 issueKey: {resp.text[:500]}")
        
        # 验证创建结果
        if verify_create:
            try:
                # 调用get_issue_detail验证工单是否真实存在
                _, created_detail = self.get_issue_detail(key)
                if created_detail and created_detail.get('key') == key:
                    print(f"[创建验证] 工单 {key} 已成功创建，标题: {created_detail.get('summary')}")
                else:
                    raise Exception(f"创建验证失败：工单 {key} 不存在或信息不匹配")
            except Exception as e:
                raise Exception(f"创建验证失败: {str(e)}")
                
        return key

    def _rest_value_for_field(self, fid, meta, value):
        """把字段值转换为 REST PUT 所需的 JSON 结构（选项型用 {"id":...}）。"""
        schema = meta.get("schema", {}) or {}
        ftype = schema.get("type")
        allowed = meta.get("allowedValues")
        is_cascading = (ftype in ("cascading", "option-with-child")
                        or (bool(allowed) and isinstance(allowed[0], dict) and bool(allowed[0].get("children"))))
        if allowed and not is_cascading:
            if isinstance(value, list):
                return [{"id": self._find_option_id(allowed, v)} for v in value]
            return {"id": self._find_option_id(allowed, value)}
        if is_cascading:
            raise NotImplementedError(f"级联字段「{meta.get('name')}」暂不支持通过 update_issue 编辑")
        return value

    def update_issue(self, issue_key, description=None, summary=None, assignee=None,
                     fields=None, project_key="IKQDM551EA", issuetype_name="ST-BUG"):
        """
        更新已有 JIRA 问题（REST PUT /rest/api/2/issue/{key}）。
        用于创建后修正内容（如把描述改成历史版式）。

        :param issue_key: 要更新的问题 Key，如 'IKQDM551EA-647'
        :param description: 描述，映射到 customfield_10244（BUG描述）
        :param summary: 标题
        :param assignee: 处理人邮箱（如 'venus.li@ikotek.com'）
        :param fields: 其它自定义字段 dict；key 可为 customfield_XXXXX 或中文名称；
                       选项型值可为选项 ID 或 label（自动解析为 {"id":...}）。
        :param project_key / issuetype_name: 仅用于解析 fields 中的选项 ID（默认 IKQDM551EA / ST-BUG）
        :return: True 表示有字段被更新
        """
        if not self.is_logged_in:
            if not self.login():
                raise Exception("JIRA登录失败")
        update = {}
        if summary is not None:
            update["summary"] = summary
        if description is not None:
            update["customfield_10244"] = description
        if assignee is not None:
            update["assignee"] = {"name": assignee}
        if fields:
            meta = self.get_create_meta(project_key=project_key, issuetype_name=issuetype_name)
            fields_meta = meta["fields"]
            name2id = {}
            for fid, m in fields_meta.items():
                if fid.startswith("customfield_") and m.get("name"):
                    name2id[m["name"]] = fid
            for key, val in fields.items():
                if val is None:
                    continue
                fid = key if str(key).startswith("customfield_") else name2id.get(key)
                if not fid or fid not in fields_meta:
                    print(f"WARNING 未知字段，已忽略: {key}")
                    continue
                if fid == "customfield_10244" and description is not None:
                    continue
                update[fid] = self._rest_value_for_field(fid, fields_meta[fid], val)
        if not update:
            return False
        resp = self.session.put(
            f"{self.base_url}/rest/api/2/issue/{issue_key}",
            json={"fields": update},
            headers={"Content-Type": "application/json", "Accept": "application/json"})
        if resp.status_code not in (200, 204):
            raise Exception(f"更新JIRA失败，状态码: {resp.status_code} - {resp.text[:500]}")
        return True

    def get_issue_raw_fields(self, issue_key):
        """获取某个历史问题的原始 fields（含自定义字段的选项 ID），用于借鉴历史值。"""
        if not self.is_logged_in:
            if not self.login():
                raise Exception("JIRA登录失败")
        resp = self.session.get(
            f"{self.base_url}/rest/api/2/issue/{issue_key}",
            headers={"Accept": "application/json",
                     "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"})
        if resp.status_code != 200:
            raise Exception(f"获取问题 {issue_key} 详情失败: {resp.status_code}")
        return resp.json().get("fields", {})

    # ------------------------------------------------------------------
    # 完整数据集获取（先取全部JIRA列表，再逐个取详细数据）
    # 思路：
    #   1) search_issues()  —— 用 REST /rest/api/2/search 分页拿到全部问题 Key（轻量、快）
    #   2) fetch_complete_dataset() —— 对每个 Key 调 get_issue_detail() 取完整详情，
    #      并发拉取，最终汇成一份「完整数据集」(list[detail])，可落盘为 JSON 供后续复用。
    # ------------------------------------------------------------------
    def search_issues(self, jql, max_results=None, fields=None, batch_size=100, exclude_closed=True):
        """
        通过 JQL 分页获取全部 JIRA 问题（仅基础字段：key/summary/status/priority/issuetype）
        使用 JIRA 标准 REST 接口 /rest/api/2/search，比 HTML 抓取更稳、可分页
        :param jql: JQL 查询语句
        :param max_results: 最多返回条数（None=不限制）
        :param fields: 需要的字段列表（默认返回基础字段）
        :param batch_size: 每页大小（默认100，最大100）
        :param exclude_closed: 是否默认排除 ST_Closed（当 JQL 未显式写 status 时）
        :return: list[dict]，每项含 key/summary/status/priority/issuetype
        """
        if not self.is_logged_in:
            if not self.login():
                raise Exception("JIRA登录失败")

        jql, _ = normalize_jql(jql, exclude_closed=exclude_closed)

        if fields is None:
            fields = ["summary", "status", "priority", "issuetype"]

        all_issues = []
        start_at = 0
        while True:
            if max_results is not None and len(all_issues) >= max_results:
                break
            url = f"{self.base_url}/rest/api/2/search"
            params = {
                "jql": jql,
                "startAt": start_at,
                "maxResults": min(batch_size, max_results - len(all_issues)) if max_results is not None else batch_size,
                "fields": ",".join(fields),
            }
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/127.0.0.0 Safari/537.36',
                'Accept': 'application/json'
            }
            resp = self.session.get(url, params=params, headers=headers)
            if resp.status_code != 200:
                raise Exception(f"JQL搜索失败，状态码: {resp.status_code} - {resp.text[:200]}")
            data = resp.json()
            issues = data.get("issues", [])
            if not issues:
                break
            for it in issues:
                f = it.get("fields", {})
                all_issues.append({
                    "key": it.get("key", ""),
                    "summary": f.get("summary", ""),
                    "status": self._name_of(f.get("status")),
                    "priority": self._name_of(f.get("priority")),
                    "issuetype": self._name_of(f.get("issuetype")),
                })
            total = data.get("total", 0)
            start_at += len(issues)
            if start_at >= total or len(issues) < batch_size:
                break

        if max_results is not None:
            all_issues = all_issues[:max_results]
        return all_issues

    def fetch_complete_dataset(self, jql, max_results=None, fields=None,
                               concurrency=4, progress=None, exclude_closed=True):
        """
        获取「完整数据集」：先用 search_issues 拿到全部 Key，再并发获取每个问题的完整详情。
        :param jql: JQL 查询语句
        :param max_results: 最多处理的问题数（None=全部）
        :param fields: 传给 get_issue_detail 的字段过滤（None=全部字段）
        :param concurrency: 并发拉取详情的线程数（默认4）
        :param progress: 进度回调 progress(done, total)
        :param exclude_closed: 是否默认排除 ST_Closed（当 JQL 未显式写 status 时）
        :return: (details, errors)
                 details: list[dict]，每个问题解析后的完整 detail（顺序与搜索一致）
                 errors: dict{key: 错误信息}，仅含拉取失败的 Key
        """
        basics = self.search_issues(jql, max_results=max_results, fields=fields, exclude_closed=exclude_closed)
        keys = [b["key"] for b in basics]
        n = len(keys)
        details = [None] * n
        errors = {}

        # 每个线程持有一个独立 JIRAClient（requests.Session 非线程安全），各自登录一次
        tlocal = threading.local()

        def worker(idx, key):
            try:
                if getattr(tlocal, "client", None) is None:
                    tlocal.client = JIRAClient(self.account, self.password, self.base_url)
                    tlocal.client.login()
                raw, detail = tlocal.client.get_issue_detail(key, fields=fields)
                return idx, detail, None
            except Exception as e:  # 单条失败不影响整体
                return idx, None, str(e)

        with ThreadPoolExecutor(max_workers=max(1, min(concurrency, n or 1))) as ex:
            futures = [ex.submit(worker, i, k) for i, k in enumerate(keys)]
            done = 0
            for fut in as_completed(futures):
                idx, detail, err = fut.result()
                done += 1
                if err:
                    errors[keys[idx]] = err
                    # 用基础信息兜底，保证数据集连续性
                    b = basics[idx]
                    details[idx] = {
                        **b,
                        "description": "", "components": [], "fixVersions": [],
                        "labels": [], "custom_fields": {}, "comments": [],
                        "attachments": [], "issuelinks": [],
                        "_fetch_error": err,
                    }
                else:
                    details[idx] = detail
                if progress:
                    progress(done, n)

        return details, errors

class JIRACompleteDataset:
    """
    完整数据集的存取与基础分析（基于已拉取好的 details，无需再访问 JIRA）。
    这样「取数据」与「用数据」解耦：先把完整数据落盘为 JSON，后续报告/过滤/统计都直接读 JSON。
    """

    # 严重等级 / 项目 字段的中文 label 候选（不同 JIRA 配置可能略有差异，按包含匹配兜底）
    SEVERITY_LABEL_CANDIDATES = ["BUG严重等级", "严重等级", "BUG严重程度", "Severity", "严重程度"]
    PROJECT_LABEL_CANDIDATES = ["BUG发现的项目", "发现项目", "Project", "项目"]

    @staticmethod
    def save_json(details, path):
        """将完整数据集落盘为 JSON，便于后续复用"""
        with open(path, "w", encoding="utf-8") as f:
            json.dump(details, f, ensure_ascii=False, indent=2)
        return os.path.abspath(path)

    @staticmethod
    def load_json(path):
        """从 JSON 读取完整数据集"""
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)

    @staticmethod
    def _match_label(custom_fields, candidates):
        for cand in candidates:
            if cand in custom_fields:
                return custom_fields[cand]
        # 兜底：按包含关系模糊匹配
        for cand in candidates:
            for k in custom_fields:
                if cand in k:
                    return custom_fields[k]
        return ""

    @classmethod
    def analyze(cls, details):
        """
        基于完整数据集做统计（不访问 JIRA）
        :return: dict{total, by_status, by_severity, by_project, by_type, fetch_errors}
        """
        by_status, by_severity, by_project, by_type = {}, {}, {}, {}
        fetch_errors = 0
        for d in details:
            if d.get("_fetch_error"):
                fetch_errors += 1
            st = d.get("status") or "未知"
            by_status[st] = by_status.get(st, 0) + 1
            sev = cls._match_label(d.get("custom_fields", {}), cls.SEVERITY_LABEL_CANDIDATES) or "未知"
            by_severity[sev] = by_severity.get(sev, 0) + 1
            proj = cls._match_label(d.get("custom_fields", {}), cls.PROJECT_LABEL_CANDIDATES) or "未知"
            by_project[proj] = by_project.get(proj, 0) + 1
            it = d.get("issuetype") or "未知"
            by_type[it] = by_type.get(it, 0) + 1
        return {
            "total": len(details),
            "fetch_errors": fetch_errors,
            "by_status": by_status,
            "by_severity": by_severity,
            "by_project": by_project,
            "by_type": by_type,
        }

class JIRADataProcessor:
    """
    JIRA数据处理器，提供过滤、排序、统计等功能
    """
    
    # 严重等级排序权重
    LEVEL_PRIORITY = {
        "A-Critical": 1,
        "B-Major": 2,
        "C-Normal": 3,
        "D-Minor": 4,
        "E-Enhancement": 5,
        "F-Trivial": 6
    }
    
    def __init__(self, data):
        """
        初始化数据处理器
        :param data: JIRA原始数据列表
        """
        self.raw_data = data
        self.processed_data = data.copy()
    
    def filter_closed(self, closed_status="ST_Closed"):
        """
        过滤掉已关闭的缺陷
        :param closed_status: 已关闭状态的标识
        :return: 过滤后的数据
        """
        self.processed_data = [
            record for record in self.processed_data 
            if record.get("Bug状态", "").strip() != closed_status
        ]
        return self.processed_data
    
    def filter_by_level(self, min_level="C-Normal"):
        """
        按严重等级过滤，只保留大于等于指定等级的缺陷
        :param min_level: 最小严重等级
        :return: 过滤后的数据
        """
        min_priority = self.LEVEL_PRIORITY.get(min_level, 99)
        self.processed_data = [
            record for record in self.processed_data
            if self.LEVEL_PRIORITY.get(record.get("Bug严重等级", "").strip(), 99) <= min_priority
        ]
        return self.processed_data
    
    def filter_by_project(self, project_name):
        """
        按项目名称过滤
        :param project_name: 项目名称
        :return: 过滤后的数据
        """
        self.processed_data = [
            record for record in self.processed_data
            if project_name in record.get("Bug发现项目", "")
        ]
        return self.processed_data
    
    def sort_by_level(self, reverse=False):
        """
        按严重等级排序
        :param reverse: 是否倒序
        :return: 排序后的数据
        """
        def get_level_priority(record):
            level = record.get("Bug严重等级", "").strip()
            return self.LEVEL_PRIORITY.get(level, 99)
        
        self.processed_data.sort(key=get_level_priority, reverse=reverse)
        return self.processed_data
    
    def sort_by_status(self):
        """
        按状态排序
        :return: 排序后的数据
        """
        status_order = {"Open": 1, "In Progress": 2, "Resolved": 3, "Closed": 4}
        self.processed_data.sort(key=lambda x: status_order.get(x.get("Bug状态", ""), 99))
        return self.processed_data
    
    def get_statistics(self):
        """
        获取数据统计信息
        :return: 统计字典
        """
        total = len(self.processed_data)
        level_stats = {}
        status_stats = {}
        project_stats = {}
        
        for record in self.processed_data:
            # 按等级统计
            level = record.get("Bug严重等级", "未知")
            level_stats[level] = level_stats.get(level, 0) + 1
            
            # 按状态统计
            status = record.get("Bug状态", "未知")
            status_stats[status] = status_stats.get(status, 0) + 1
            
            # 按项目统计
            project = record.get("Bug发现项目", "未知")
            project_stats[project] = project_stats.get(project, 0) + 1
        
        return {
            "total_count": total,
            "by_level": level_stats,
            "by_status": status_stats,
            "by_project": project_stats
        }
    
    def get_processed_data(self):
        """获取处理后的数据"""
        return self.processed_data

class DocxReportGenerator:
    """
    DOCX报告生成器
    """
    
    def __init__(self, title="JIRA缺陷统计报告"):
        """
        初始化报告生成器
        :param title: 报告标题
        """
        if not DOCX_SUPPORTED:
            raise Exception("DOCX功能不可用，请安装python-docx库")
        
        self.title = title
        self.doc = Document()
    
    def add_title(self, title_text):
        """添加文档标题"""
        title_para = self.doc.add_heading(title_text, level=1)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in title_para.runs:
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    def add_info_section(self, version="", total_count=0, statistics=None):
        """
        添加信息区域
        :param version: 版本号
        :param total_count: 缺陷总数
        :param statistics: 统计信息字典
        """
        # 基本信息
        info_para = self.doc.add_paragraph()
        info_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        if version:
            run1 = info_para.add_run(f"版本号: ")
            run1.font.size = Pt(11)
            run1.font.bold = True
            run2 = info_para.add_run(f"{version}\n")
            run2.font.size = Pt(11)
        
        run3 = info_para.add_run(f"缺陷总数: ")
        run3.font.size = Pt(11)
        run3.font.bold = True
        run4 = info_para.add_run(f"{total_count}\n")
        run4.font.size = Pt(11)
        run4.font.color.rgb = RGBColor(255, 0, 0)  # 红色
        
        # 统计信息
        if statistics:
            info_para.add_run("\n")
            
            # 按等级统计
            run5 = info_para.add_run(f"按严重等级统计:\n")
            run5.font.size = Pt(11)
            run5.font.bold = True
            
            for level, count in statistics.get("by_level", {}).items():
                info_para.add_run(f"  {level}: {count}个\n")
            
            info_para.add_run("\n")
            
            # 按状态统计
            run6 = info_para.add_run(f"按状态统计:\n")
            run6.font.size = Pt(11)
            run6.font.bold = True
            
            for status, count in statistics.get("by_status", {}).items():
                info_para.add_run(f"  {status}: {count}个\n")
    
    def create_table(self, headers, data):
        """
        创建缺陷表格
        :param headers: 表头列表
        :param data: 数据列表
        """
        if not data:
            self.doc.add_paragraph("暂无数据")
            return
        
        # 创建表格
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置表头
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # 表头样式
            for paragraph in hdr_cells[i].paragraphs:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.name = 'Arial'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # 填充数据
        left_idx = headers.index("Bug描述") if "Bug描述" in headers else -1
        for row_data in data:
            row_cells = table.add_row().cells
            for i, cell_value in enumerate(row_data):
                row_cells[i].text = str(cell_value)
                row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                
                # 数据行样式
                for paragraph in row_cells[i].paragraphs:
                    # Bug描述列左对齐，其他居中
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT if i == left_idx else WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        
        # 设置列宽
        default_widths = [Inches(1.0), Inches(1.0), Inches(3.5), Inches(2.5), Inches(1.5)]
        widths = default_widths[:len(headers)]
        for row in table.rows:
            for idx, width in enumerate(widths):
                if idx < len(row.cells):
                    row.cells[idx].width = width
    
    def add_heading2(self, text):
        """添加二级标题"""
        h = self.doc.add_heading(text, level=2)
        for run in h.runs:
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.name = 'Arial'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    def add_paragraph_text(self, text, bold=False, size=11):
        """添加普通段落文本"""
        p = self.doc.add_paragraph()
        r = p.add_run(text if text else '（无）')
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.name = 'Arial'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        return p
    
    def add_kv(self, key, value):
        """添加 字段: 值 形式的段落"""
        p = self.doc.add_paragraph()
        r1 = p.add_run(f"{key}: ")
        r1.font.bold = True
        r1.font.size = Pt(11)
        r1.font.name = 'Arial'
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r2 = p.add_run(str(value))
        r2.font.size = Pt(11)
        r2.font.name = 'Arial'
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        return p
    
    def create_kv_table(self, kv_dict):
        """创建 键-值 两列表格"""
        if not kv_dict:
            self.doc.add_paragraph('（无）')
            return
        table = self.doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for k, v in kv_dict.items():
            row_cells = table.add_row().cells
            row_cells[0].text = str(k)
            row_cells[1].text = str(v)
            for cidx in (0, 1):
                for paragraph in row_cells[cidx].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Arial'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    def save(self, filename=None):
        """
        保存文档
        :param filename: 保存路径，默认自动生成
        :return: 保存的文件路径
        """
        if not filename:
            filename = f"JIRA缺陷报告_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        self.doc.save(filename)
        file_path = os.path.abspath(filename)
        return file_path

# 快捷函数
def generate_markdown(account, password, jql_query=None, jira_url=None, version="", output_file=None, exclude_closed=True, with_comments=True, remark_map=None):
    """
    快捷生成JIRA缺陷数据Markdown格式
    :param account: JIRA账号
    :param password: JIRA密码
    :param jql_query: JQL查询语句
    :param jira_url: JIRA搜索URL（与jql_query二选一）
    :param version: 版本号
    :param output_file: 输出文件路径，默认返回Markdown字符串
    :param exclude_closed: 是否默认排除 ST_Closed（当 JQL 未显式写 status 时）
    :param with_comments: 是否拉取并在「备注」列写入原始 JIRA 评论
    :param remark_map: 外部提供的备注映射 {Bug号: 备注文本}（优先级高于 with_comments）
    :return: Markdown内容或文件路径
    """
    # 1. 提取数据
    client = JIRAClient(account, password)
    
    if jql_query:
        headers, raw_data = client.query_by_jql(jql_query, exclude_closed=exclude_closed)
    elif jira_url:
        headers, raw_data = client.query_by_url(jira_url, exclude_closed=exclude_closed)
    else:
        raise Exception("必须提供jql_query或jira_url参数")
    
    print(f"提取到 {len(raw_data)} 条记录")
    
    # 2. 处理数据
    processor = JIRADataProcessor(raw_data)
    if exclude_closed:
        processor.filter_closed()
    processor.sort_by_level()
    processed_data = processor.get_processed_data()
    statistics = processor.get_statistics()
    
    print(f"过滤后剩余 {len(processed_data)} 条有效记录")

    # 4. 拉取每条记录的 JIRA 备注/评论
    if remark_map is None and with_comments:
        print("开始拉取每条问题的备注...")
        details_map = _fetch_details_map(client, processed_data)
        remark_map = {k: _format_comments(_comments_of(v), sep="<br>", newline_repl="<br>")
                      for k, v in details_map.items()}
    else:
        remark_map = remark_map or {}

    # 5. 生成Markdown
    md_content = []
    md_content.append("# JIRA缺陷统计报告\n")
    
    if version:
        md_content.append(f"**版本号**: {version}\n")
    md_content.append(f"**缺陷总数**: {len(processed_data)} 个\n\n")
    
    # 按等级统计
    md_content.append("## 按严重等级统计\n")
    for level, count in statistics.get("by_level", {}).items():
        md_content.append(f"- {level}: {count}个\n")
    md_content.append("\n")
    
    # 按状态统计
    md_content.append("## 按状态统计\n")
    for status, count in statistics.get("by_status", {}).items():
        md_content.append(f"- {status}: {count}个\n")
    md_content.append("\n")
    
    # 缺陷详情表格（含备注列）
    md_content.append("## 缺陷详情\n")
    md_content.append("| Bug号 | 严重等级 | Bug描述 | 备注 |\n")
    md_content.append("|-------|----------|---------|------|\n")
    for record in processed_data:
        bug_id = record.get("Bug号", "")
        level = record.get("Bug严重等级", "")
        desc = record.get("Bug描述", "").replace('\n', ' ').replace('|', '\\|')
        remark = remark_map.get(bug_id, "")
        md_content.append(f"| {bug_id} | {level} | {desc} | {remark} |\n")
    
    markdown = ''.join(md_content)
    
    if output_file:
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(markdown)
        print(f"Markdown文件生成成功：{output_file}")
        return output_file
    else:
        return markdown

def generate_report(account, password, jql_query=None, jira_url=None, version="", output_file=None, output_format="markdown", exclude_closed=True, with_comments=True, remark_map=None):
    """
    快捷生成JIRA缺陷报告
    :param account: JIRA账号
    :param password: JIRA密码
    :param jql_query: JQL查询语句
    :param jira_url: JIRA搜索URL（与jql_query二选一）
    :param version: 版本号
    :param output_file: 输出文件路径
    :param output_format: 输出格式，支持'docx'或'markdown'，默认'docx'
    :param exclude_closed: 是否默认排除 ST_Closed（当 JQL 未显式写 status 时）
    :param with_comments: 是否拉取并在「备注」列写入原始 JIRA 评论
    :param remark_map: 外部提供的备注映射 {Bug号: 备注文本}（优先级高于 with_comments）
    :return: 生成的报告路径
    """
    if output_format.lower() == 'markdown':
        return generate_markdown(account, password, jql_query, jira_url, version, output_file,
                                  exclude_closed=exclude_closed, with_comments=with_comments, remark_map=remark_map)
    
    # 1. 提取数据
    client = JIRAClient(account, password)
    
    if jql_query:
        headers, raw_data = client.query_by_jql(jql_query, exclude_closed=exclude_closed)
    elif jira_url:
        headers, raw_data = client.query_by_url(jira_url, exclude_closed=exclude_closed)
    else:
        raise Exception("必须提供jql_query或jira_url参数")
    
    print(f"提取到 {len(raw_data)} 条记录")
    
    # 2. 处理数据
    processor = JIRADataProcessor(raw_data)
    if exclude_closed:
        processor.filter_closed()
    processor.sort_by_level()
    processed_data = processor.get_processed_data()
    statistics = processor.get_statistics()
    
    print(f"过滤后剩余 {len(processed_data)} 条有效记录")
    
    # 3. 备注内容
    if remark_map is None and with_comments:
        print("开始拉取每条问题的备注...")
        details_map = _fetch_details_map(client, processed_data)
        remark_map = {k: _format_comments(_comments_of(v), sep="\n", newline_repl="\n")
                      for k, v in details_map.items()}
    else:
        remark_map = remark_map or {}
    
    # 4. 生成报告
    output_path = _build_defect_docx(processed_data, statistics, version, remark_map, output_file)
    print(f"报告生成成功：{output_path}")
    return output_path


def export_three_reports(account, password, jql, version="", output_dir=".",
                         exclude_closed=True, ai_notes=None,
                         ai_input_json="jira_ai_input.json"):
    """
    一次性输出三份报告（核心入口）：
      1) JIRA缺陷报告.docx                              —— 基础报告（备注列留空）
      2) JIRA缺陷报告_含备注.docx                        —— 备注列写入原始 JIRA 评论
      3) JIRA缺陷报告_含备注_AI处理备注生成新的备注信息.docx —— 备注列写入 AI 分析后的类人工备注
    同时输出 jira_ai_input.json 供 AI（WorkBuddy）分析生成第3份的备注。
    :param ai_notes: dict{Bug号: AI生成的备注文本}。为 None 时第3份报告备注列留空（占位）。
    :param ai_input_json: 供 AI 分析用的输入 JSON 文件名。
    :return: dict{report1, report2, report3, ai_input, count}
    """
    os.makedirs(output_dir, exist_ok=True)
    client = JIRAClient(account, password)
    headers, raw_data = client.query_by_jql(jql, exclude_closed=exclude_closed)
    print(f"提取到 {len(raw_data)} 条记录")

    processor = JIRADataProcessor(raw_data)
    if exclude_closed:
        processor.filter_closed()
    processor.sort_by_level()
    processed_data = processor.get_processed_data()
    statistics = processor.get_statistics()
    print(f"缺陷总数：{len(processed_data)}")

    # 拉取完整详情（含评论），一次复用
    print("开始拉取每条问题的详情/评论...")
    details_map = _fetch_details_map(client, processed_data)
    raw_remark_map = {k: _format_comments(_comments_of(v), sep="\n", newline_repl="\n")
                      for k, v in details_map.items()}

    # 报告1：基础（备注空）
    r1 = _build_defect_docx(processed_data, statistics, version, {},
                            os.path.join(output_dir, "JIRA缺陷报告.docx"))
    # 报告2：原始评论
    r2 = _build_defect_docx(processed_data, statistics, version, raw_remark_map,
                            os.path.join(output_dir, "JIRA缺陷报告_含备注.docx"))
    # 报告3：AI 备注
    ai_notes = ai_notes or {}
    ai_remark_map = {record.get("Bug号", ""): ai_notes.get(record.get("Bug号", ""), "")
                     for record in processed_data}
    r3 = _build_defect_docx(processed_data, statistics, version, ai_remark_map,
                            os.path.join(output_dir, "JIRA缺陷报告_含备注_AI处理备注生成新的备注信息.docx"))

    # AI 输入 JSON
    ai_input = _build_ai_input(processed_data, details_map)
    ai_path = os.path.join(output_dir, ai_input_json)
    with open(ai_path, "w", encoding="utf-8") as f:
        json.dump(ai_input, f, ensure_ascii=False, indent=2)
    print(f"AI 输入已生成：{ai_path}")

    return {"report1": r1, "report2": r2, "report3": r3,
            "ai_input": ai_path, "count": len(processed_data)}

# ------------------------------------------------------------------
# 单个JIRA问题详情报告生成
# ------------------------------------------------------------------
def _format_comments(comments, max_per_body=200, max_comments=5, sep="\n", newline_repl=" "):
    """
    将评论列表格式化为字符串（用于表格备注列）。
    :param comments: list[dict] 每个含 author/created/body
    :param max_per_body: 单条评论正文最多字符数（超出截断）
    :param max_comments: 最多显示几条评论
    :param sep: 多条评论之间的分隔符
    :param newline_repl: 评论正文内换行替换成什么（DOCX 用 \n，Markdown 用 <br>）
    :return: 格式化字符串；无评论返回空串
    """
    if not comments:
        return ""
    parts = []
    for i, c in enumerate(comments[:max_comments], 1):
        header = f"{i}. {c.get('author', '')} ({c.get('created', '')})"
        body = (c.get('body', '') or '').replace('\n', newline_repl).replace('\r', '')
        if len(body) > max_per_body:
            body = body[:max_per_body] + "…"
        parts.append(f"{header}：{body}")
    if len(comments) > max_comments:
        parts.append(f"...（共 {len(comments)} 条评论，仅显示前 {max_comments} 条）")
    return sep.join(parts)


def _fetch_details_map(client, records, concurrency=4):
    """
    并发获取每条记录对应的 JIRA 完整详情（含评论/描述/自定义字段）。
    :param client: JIRAClient 实例
    :param records: 含 Bug号 字段的原始/处理后记录列表
    :param concurrency: 并发线程数
    :return: dict{Bug号: detail 或 None}
    """
    def fetch_one(record):
        key = str(record.get("Bug号", "")).strip()
        if not key:
            return key, None
        try:
            raw, detail = client.get_issue_detail(key)
            return key, detail
        except Exception as e:
            print(f"  获取 {key} 详情失败：{e}")
            return key, None

    results = {}
    with ThreadPoolExecutor(max_workers=concurrency) as exe:
        futures = {exe.submit(fetch_one, r): r for r in records}
        for future in as_completed(futures):
            key, detail = future.result()
            results[key] = detail
    return results


def _comments_of(detail):
    """从 detail 中安全取评论列表"""
    if not detail:
        return []
    return detail.get("comments", []) or []


def _build_defect_docx(processed_data, statistics, version, remark_map, output_file):
    """
    构建缺陷 DOCX 报告（含 Bug号/严重等级/Bug描述/备注 四列）。
    :param remark_map: dict{Bug号: 备注文本}（备注列内容，可为空）
    :return: 保存路径
    """
    generator = DocxReportGenerator()
    generator.add_title(generator.title)
    generator.add_info_section(version, len(processed_data), statistics)
    generator.doc.add_paragraph()  # 空行

    docx_headers = ["Bug号", "Bug严重等级", "Bug描述", "备注"]
    docx_data = []
    for record in processed_data:
        bug_id = record.get("Bug号", "")
        docx_data.append([
            bug_id,
            record.get("Bug严重等级", ""),
            record.get("Bug描述", ""),
            remark_map.get(bug_id, "")
        ])

    generator.create_table(docx_headers, docx_data)
    return generator.save(output_file)


def _build_ai_input(processed_data, details_map, max_comments=20):
    """
    构建供 AI 分析用的输入 JSON 结构（每个问题含关键信息 + 评论）。
    :param processed_data: 处理后的列表记录
    :param details_map: dict{Bug号: detail}
    :return: list[dict]
    """
    items = []
    for record in processed_data:
        bug_id = record.get("Bug号", "")
        detail = details_map.get(bug_id) or {}
        # 关键自定义字段（过滤系统噪声由 get_issue_detail 已处理）
        cf = detail.get("custom_fields", {}) or {}
        # 选取对分析有用的字段
        useful = {}
        for label, val in cf.items():
            sval = str(val)
            if len(sval) > 300:
                sval = sval[:300] + "…"
            useful[label] = sval
        items.append({
            "key": bug_id,
            "severity": record.get("Bug严重等级", ""),
            "status": record.get("Bug状态", detail.get("status", "")),
            "summary": record.get("Bug描述", detail.get("summary", "")),
            "description": (detail.get("description", "") or "")[:1000],
            "custom_fields": useful,
            "comments": [
                {"author": c.get("author", ""), "created": c.get("created", ""), "body": c.get("body", "")}
                for c in _comments_of(detail)[:max_comments]
            ],
        })
    return items


def _render_issue_markdown(detail):
    """
    将一个已解析的 issue detail 渲染为 Markdown 字符串（不访问 JIRA）。
    供 generate_issue_markdown 与完整数据集报告复用。
    """
    md = []
    md.append(f"# JIRA问题详情 - {detail['key']}\n\n")
    md.append(f"**标题**: {detail['summary']}\n\n")
    md.append(f"**状态**: {detail['status']}  |  **优先级**: {detail['priority']}  |  **类型**: {detail['issuetype']}\n\n")
    md.append(f"**报告人**: {detail['reporter']}  |  **处理人**: {detail['assignee']}\n\n")
    md.append(f"**创建时间**: {detail['created']}  |  **更新时间**: {detail['updated']}\n\n")
    if detail['components']:
        md.append(f"**模块**: {', '.join(detail['components'])}\n\n")
    if detail['fixVersions']:
        md.append(f"**修复版本**: {', '.join(detail['fixVersions'])}\n\n")
    if detail['labels']:
        md.append(f"**标签**: {', '.join(detail['labels'])}\n\n")

    md.append("## 描述\n\n")
    md.append((detail['description'] or '（无）') + "\n\n")

    md.append("## 自定义字段\n\n")
    if detail['custom_fields']:
        md.append("| 字段 | 值 |\n|------|-----|\n")
        for k, v in detail['custom_fields'].items():
            cell = str(v).replace('\n', ' ').replace('|', '\\|')
            if len(cell) > 300:
                cell = cell[:300] + '…(已截断)'
            md.append(f"| {k} | {cell} |\n")
    else:
        md.append("（无）\n\n")

    if detail['comments']:
        md.append("\n## 评论\n\n")
        for c in detail['comments']:
            md.append(f"**{c['author']}** ({c['created']}):\n\n{c['body']}\n\n")

    if detail['attachments']:
        md.append("\n## 附件\n\n")
        for a in detail['attachments']:
            md.append(f"- [{a['filename']}]({a['url']}) ({a['created']})\n")

    if detail['issuelinks']:
        md.append("\n## 关联问题\n\n")
        for l in detail['issuelinks']:
            md.append(f"- {l['type']} ({l['direction']}): {l['key']} {l['summary']}\n")

    if detail.get('_fetch_error'):
        md.append(f"\n> ⚠️  该问题详情拉取失败：{detail['_fetch_error']}\n")

    return ''.join(md)

def generate_issue_markdown(account, password, issue_key, output_file=None):
    """
    生成单个JIRA问题详情的Markdown报告
    :param account: JIRA账号
    :param password: JIRA密码
    :param issue_key: 问题Key，如 'QDM565EA-396'
    :param output_file: 输出文件路径，默认返回Markdown字符串
    :return: Markdown内容或文件路径
    """
    client = JIRAClient(account, password)
    raw, detail = client.get_issue_detail(issue_key)
    content = _render_issue_markdown(detail)
    if output_file:
        with open(output_file, 'w', encoding='utf-8') as fp:
            fp.write(content)
        print(f"Markdown报告已生成：{output_file}")
        return output_file
    return content

def generate_issue_docx(account, password, issue_key, output_file=None):
    """
    生成单个JIRA问题详情的DOCX报告
    :param account: JIRA账号
    :param password: JIRA密码
    :param issue_key: 问题Key，如 'QDM565EA-396'
    :param output_file: 输出文件路径
    :return: 生成的报告路径
    """
    if not DOCX_SUPPORTED:
        raise Exception("DOCX功能不可用，请安装python-docx库")
    
    client = JIRAClient(account, password)
    raw, detail = client.get_issue_detail(issue_key)
    
    gen = DocxReportGenerator(title=f"JIRA问题详情 - {detail['key']}")
    gen.add_title(gen.title)
    gen.add_kv("标题", detail['summary'])
    gen.add_kv("状态", detail['status'])
    gen.add_kv("优先级", detail['priority'])
    gen.add_kv("类型", detail['issuetype'])
    gen.add_kv("报告人", detail['reporter'])
    gen.add_kv("处理人", detail['assignee'])
    gen.add_kv("创建时间", detail['created'])
    gen.add_kv("更新时间", detail['updated'])
    if detail['components']:
        gen.add_kv("模块", ', '.join(detail['components']))
    if detail['fixVersions']:
        gen.add_kv("修复版本", ', '.join(detail['fixVersions']))
    if detail['labels']:
        gen.add_kv("标签", ', '.join(detail['labels']))
    gen.doc.add_paragraph()
    
    gen.add_heading2("描述")
    gen.add_paragraph_text(detail['description'])
    
    gen.add_heading2("自定义字段")
    gen.create_kv_table(detail['custom_fields'])
    
    if detail['comments']:
        gen.add_heading2("评论")
        for c in detail['comments']:
            gen.add_paragraph_text(f"{c['author']} ({c['created']})：", bold=True, size=10)
            gen.add_paragraph_text(c['body'], size=10)
    
    if detail['attachments']:
        gen.add_heading2("附件")
        for a in detail['attachments']:
            gen.add_paragraph_text(f"- {a['filename']} ({a['created']})  {a['url']}", size=10)
    
    if detail['issuelinks']:
        gen.add_heading2("关联问题")
        for l in detail['issuelinks']:
            gen.add_paragraph_text(f"- {l['type']} ({l['direction']}): {l['key']} {l['summary']}", size=10)
    
    out = gen.save(output_file)
    print(f"DOCX报告已生成：{out}")
    return out

def generate_issue_report(account, password, issue_key, output_file=None, output_format="markdown"):
    """
    生成单个JIRA问题详情报告（快捷入口）
    :param account: JIRA账号
    :param password: JIRA密码
    :param issue_key: 问题Key，如 'QDM565EA-396'
    :param output_file: 输出文件路径
    :param output_format: 输出格式，支持 'docx' 或 'markdown'，默认 'markdown'
    :return: 生成的报告路径或Markdown字符串
    """
    if output_format.lower() == 'docx':
        return generate_issue_docx(account, password, issue_key, output_file)
    return generate_issue_markdown(account, password, issue_key, output_file)

# ------------------------------------------------------------------
# 完整数据集报告生成（直接基于已拉取的 details，无需再访问 JIRA）
# ------------------------------------------------------------------
def _render_complete_markdown(details, version="", include_full_details=False):
    """将一个完整数据集渲染为 Markdown 字符串"""
    stats = JIRACompleteDataset.analyze(details)
    md = []
    md.append("# JIRA 完整数据集报告\n\n")
    if version:
        md.append(f"**版本号**: {version}\n\n")
    md.append(f"**问题总数**: {stats['total']} 个")
    if stats['fetch_errors']:
        md.append(f"（其中 {stats['fetch_errors']} 个详情拉取失败，已用基础信息兜底）")
    md.append("\n\n")

    md.append("## 统计概览\n")
    md.append("\n### 按状态\n")
    for k, v in stats['by_status'].items():
        md.append(f"- {k}: {v}\n")
    md.append("\n### 按严重等级\n")
    for k, v in stats['by_severity'].items():
        md.append(f"- {k}: {v}\n")
    md.append("\n### 按发现项目\n")
    for k, v in stats['by_project'].items():
        md.append(f"- {k}: {v}\n")
    md.append("\n### 按类型\n")
    for k, v in stats['by_type'].items():
        md.append(f"- {k}: {v}\n")

    md.append("\n## 问题清单\n")
    md.append("| Key | 类型 | 状态 | 严重等级 | 处理人 | 报告人 | 标题 | 备注 |\n")
    md.append("|-----|------|------|----------|--------|--------|------|------|\n")
    cf = JIRACompleteDataset.SEVERITY_LABEL_CANDIDATES
    for d in details:
        sev = JIRACompleteDataset._match_label(d.get('custom_fields', {}), cf) or ''
        comments = d.get('comments', [])
        remark = _format_comments(comments, sep="<br>", newline_repl="<br>")
        row = [
            d.get('key', ''),
            d.get('issuetype', ''),
            d.get('status', ''),
            sev,
            d.get('assignee', ''),
            d.get('reporter', ''),
            str(d.get('summary', '')).replace('\n', ' ').replace('|', '\\|'),
            remark,
        ]
        md.append("| " + " | ".join(str(x) for x in row) + " |\n")

    if include_full_details:
        md.append("\n---\n\n## 各问题完整详情\n")
        for d in details:
            md.append(f"\n### {d.get('key', '')}\n")
            md.append(_render_issue_markdown(d))

    return ''.join(md)


def generate_complete_markdown(details, version="", output_file=None, include_full_details=False):
    """
    基于完整数据集生成 Markdown 报告（不访问 JIRA）
    :param details: fetch_complete_dataset / JIRACompleteDataset.load_json 返回的 list[detail]
    :param version: 版本号
    :param output_file: 输出路径；不传则返回字符串
    :param include_full_details: 是否追加每个问题的完整详情
    :return: 文件路径或 Markdown 字符串
    """
    content = _render_complete_markdown(details, version=version, include_full_details=include_full_details)
    if output_file:
        with open(output_file, 'w', encoding='utf-8') as fp:
            fp.write(content)
        print(f"完整数据集 Markdown 报告已生成：{output_file}")
        return output_file
    return content


def generate_complete_docx(details, version="", output_file=None, include_full_details=False):
    """
    基于完整数据集生成 DOCX 报告（不访问 JIRA）
    :param details: list[detail]
    :param version: 版本号
    :param output_file: 输出路径
    :param include_full_details: 是否追加每个问题的完整详情
    :return: 报告路径
    """
    if not DOCX_SUPPORTED:
        raise Exception("DOCX功能不可用，请安装python-docx库")

    stats = JIRACompleteDataset.analyze(details)
    gen = DocxReportGenerator(title="JIRA 完整数据集报告")
    gen.add_title(gen.title)
    if version:
        gen.add_kv("版本号", version)
    gen.add_kv("问题总数", stats['total'])
    if stats['fetch_errors']:
        gen.add_kv("详情拉取失败", f"{stats['fetch_errors']} 个（已用基础信息兜底）")
    gen.doc.add_paragraph()

    gen.add_heading2("统计概览")
    gen.add_paragraph_text("按状态：", bold=True, size=10)
    for k, v in stats['by_status'].items():
        gen.add_paragraph_text(f"  {k}: {v}", size=10)
    gen.add_paragraph_text("按严重等级：", bold=True, size=10)
    for k, v in stats['by_severity'].items():
        gen.add_paragraph_text(f"  {k}: {v}", size=10)
    gen.add_paragraph_text("按发现项目：", bold=True, size=10)
    for k, v in stats['by_project'].items():
        gen.add_paragraph_text(f"  {k}: {v}", size=10)

    gen.add_heading2("问题清单")
    cf = JIRACompleteDataset.SEVERITY_LABEL_CANDIDATES
    headers = ["Key", "类型", "状态", "严重等级", "处理人", "报告人", "标题", "备注"]
    table_data = []
    for d in details:
        sev = JIRACompleteDataset._match_label(d.get('custom_fields', {}), cf) or ''
        comments = d.get('comments', [])
        remark = _format_comments(comments, sep="\n", newline_repl="\n")
        table_data.append([
            d.get('key', ''), d.get('issuetype', ''), d.get('status', ''),
            sev, d.get('assignee', ''), d.get('reporter', ''),
            str(d.get('summary', ''))[:80],
            remark,
        ])
    gen.create_table(headers, table_data)

    if include_full_details:
        gen.add_heading2("各问题完整详情")
        for d in details:
            gen.add_heading2(d.get('key', ''))
            gen.add_kv("标题", d.get('summary', ''))
            gen.add_kv("状态", d.get('status', ''))
            gen.add_kv("优先级", d.get('priority', ''))
            gen.add_kv("类型", d.get('issuetype', ''))
            gen.add_kv("报告人", d.get('reporter', ''))
            gen.add_kv("处理人", d.get('assignee', ''))
            gen.add_kv("创建时间", d.get('created', ''))
            gen.add_kv("更新时间", d.get('updated', ''))
            if d.get('components'):
                gen.add_kv("模块", ', '.join(d['components']))
            if d.get('fixVersions'):
                gen.add_kv("修复版本", ', '.join(d['fixVersions']))
            gen.add_heading2("描述")
            gen.add_paragraph_text(d.get('description', ''))
            gen.add_heading2("自定义字段")
            gen.create_kv_table(d.get('custom_fields', {}))
            if d.get('comments'):
                gen.add_heading2("评论")
                for c in d['comments']:
                    gen.add_paragraph_text(f"{c['author']} ({c['created']})：", bold=True, size=10)
                    gen.add_paragraph_text(c['body'], size=10)
            if d.get('attachments'):
                gen.add_heading2("附件")
                for a in d['attachments']:
                    gen.add_paragraph_text(f"- {a['filename']} ({a['created']})  {a['url']}", size=10)
            if d.get('issuelinks'):
                gen.add_heading2("关联问题")
                for l in d['issuelinks']:
                    gen.add_paragraph_text(f"- {l['type']} ({l['direction']}): {l['key']} {l['summary']}", size=10)

    out = gen.save(output_file)
    print(f"完整数据集 DOCX 报告已生成：{out}")
    return out


def generate_complete_report(details, version="", output_file=None,
                             output_format="markdown", include_full_details=False):
    """
    基于完整数据集生成报告（快捷入口，不访问 JIRA）
    :param details: list[detail]
    :param output_format: 'markdown' 或 'docx'
    :return: 报告路径或 Markdown 字符串
    """
    if output_format.lower() == 'docx':
        return generate_complete_docx(details, version, output_file, include_full_details)
    return generate_complete_markdown(details, version, output_file, include_full_details)


def export_complete_dataset(account, password, jql, output_dir=".", version="",
                            max_results=None, concurrency=4, output_format="both",
                            include_full_details=False, dataset_filename="jira_complete_dataset.json",
                            exclude_closed=True):
    """
    一站式导出「完整数据集」：
      1) 用 JQL 拉取全部问题列表；
      2) 并发获取每个问题的完整详情；
      3) 落盘为 JSON 完整数据集（供后续任意操作复用，无需再访问 JIRA）；
      4) 同时生成 Markdown / DOCX 报告。
    :param account/password: JIRA 账号密码
    :param jql: JQL 查询语句
    :param output_dir: 输出目录
    :param version: 版本号（写入报告）
    :param max_results: 最多处理的问题数（None=全部）
    :param concurrency: 并发拉取详情的线程数
    :param output_format: 'markdown' / 'docx' / 'both'
    :param include_full_details: 报告是否含每个问题完整详情
    :param dataset_filename: 完整数据集 JSON 文件名
    :param exclude_closed: 是否默认排除 ST_Closed（当 JQL 未显式写 status 时）
    :return: dict{json, markdown, docx, stats, errors}
    """
    os.makedirs(output_dir, exist_ok=True)
    client = JIRAClient(account, password)
    print(f"▶ 步骤1/2：按 JQL 拉取问题列表 ...")
    details, errors = client.fetch_complete_dataset(
        jql, max_results=max_results, concurrency=concurrency,
        exclude_closed=exclude_closed,
        progress=lambda done, total: print(f"  拉取详情 {done}/{total}", end="\r")
    )
    print(f"\n✅ 完整数据集就绪：共 {len(details)} 个问题（失败 {len(errors)} 个）")

    json_path = JIRACompleteDataset.save_json(details, os.path.join(output_dir, dataset_filename))
    print(f"📦 完整数据集已落盘：{json_path}")

    result = {"json": json_path, "markdown": None, "docx": None,
              "stats": JIRACompleteDataset.analyze(details), "errors": errors}

    fmt = output_format.lower()
    if fmt in ("markdown", "both"):
        md_path = os.path.join(output_dir, "JIRA完整数据集报告.md")
        result["markdown"] = generate_complete_markdown(
            details, version=version, output_file=md_path, include_full_details=include_full_details)
    if fmt in ("docx", "both"):
        dx_path = os.path.join(output_dir, "JIRA完整数据集报告.docx")
        result["docx"] = generate_complete_docx(
            details, version=version, output_file=dx_path, include_full_details=include_full_details)

    return result


# ==================================================================
# 历史测试 JIRA 内容样式模板
# 参考 IKQDM551EA-609 / -603 / -646 等真实测试类缺陷的写法整理。
# 历史版式要点：
#   1) 标题固定为：ST[QDM_551][OTA]<现象>；期望<期望>；<概率>
#   2) 描述固定分块：[测试环境]：/ [测试步骤]：/ 期望结果:/ [测试现象]/ [概率]：
#      （注意 [测试环境]：、[测试步骤]： 带中文全角冒号；期望结果: 仅半角冒号；
#       [测试现象] 无冒号；[概率]：再带中文全角冒号）
#   3) [测试现象] 通常直接重复标题整句。
# ==================================================================
DEFAULT_TEST_TAG = "ST[QDM_551][OTA]"


def _as_lines(x):
    if x is None:
        return []
    if isinstance(x, (list, tuple)):
        return [str(i) for i in x]
    return [str(x)]


def build_bug_summary(test_tag=DEFAULT_TEST_TAG, phenomenon="", expectation="", probability=""):
    """
    构造标题：ST[QDM_551][OTA]<现象>；期望<期望>；<概率>
    例：ST[QDM_551][OTA]挂测组合升级，升级失败；期望升级成功；大概率
    """
    return f"{test_tag}{phenomenon}；期望{expectation}；{probability}"


def build_bug_description(env=None, steps=None, expected=None, phenomenon="", probability="", summary=None):
    """
    构造描述，严格遵循历史测试 JIRA 版式（参考 IKQDM551EA-609 等）：

        [测试环境]：
        模块IMEI编号：xxx
        SIM卡：xxx
        AT通信口：/
        网络配置：xxx
        产品线：xxx
        [测试步骤]：
        1.动作
        期望结果:
        1.期望
        [测试现象]
        <现象（与标题完全一致）>
        [概率]：<概率>

    :param env: dict，键 imei / sim / at_port / network / product_line
    :param steps: 字符串或列表，测试步骤（编号自动生成）
    :param expected: 字符串或列表，期望结果
    :param phenomenon: 测试现象，若提供summary则优先使用summary作为测试现象
    :param probability: 概率词（必现 / 大概率 / 小概率 / ...）
    :param summary: 标题内容，若提供则[测试现象]字段完全等于该值
    """
    env = env or {}
    lines = ["[测试环境]："]
    lines.append(f"模块IMEI编号：{env.get('imei', '')}")
    lines.append(f"SIM卡：{env.get('sim', '')}")
    lines.append(f"AT通信口：{env.get('at_port', '/')}")
    lines.append(f"网络配置：{env.get('network', '')}")
    lines.append(f"产品线：{env.get('product_line', '')}")
    lines.append("[测试步骤]：")
    for i, s in enumerate(_as_lines(steps), 1):
        lines.append(f"{i}.{s}")
    lines.append("期望结果:")
    for i, s in enumerate(_as_lines(expected), 1):
        lines.append(f"{i}.{s}")
    lines.append("[测试现象]")
    # 若提供summary参数，则测试现象完全等于标题，否则使用phenomenon
    lines.append(summary if summary is not None else phenomenon)
    lines.append(f"[概率]：{probability}")
    return "\n".join(lines)


def create_issue_draft(account, password, project_key=None, issuetype_name="ST-BUG",
                       summary=None, description=None, fields=None, assignee=None,
                       priority=None, components=None, attachments=None, links=None,
                       project_id=None):
    """
    只读地生成创建草稿（不创建、不上传附件）。返回结构化草稿 dict，
    供「先确认再创建」流程展示给用户确认。详见 JIRAClient.build_create_draft。
    """
    client = JIRAClient(account, password)
    if not client.login():
        raise Exception("JIRA登录失败")
    return client.build_create_draft(
        project_key=project_key, issuetype_name=issuetype_name, summary=summary,
        description=description, fields=fields, assignee=assignee, priority=priority,
        components=components, attachments=attachments, links=links, project_id=project_id)


def create_issue(account, password, project_key=None, issuetype_name="ST-BUG",
                summary=None, description=None, fields=None, assignee=None,
                priority=None, components=None, attachments=None, links=None,
                project_id=None):
    """
    创建 JIRA 问题的快捷入口（封装登录 + 创建）。
    详见 JIRAClient.create_issue 的参数说明。
    :return: 新问题 Key 字符串
    """
    client = JIRAClient(account, password)
    if not client.login():
        raise Exception("JIRA登录失败")
    return client.create_issue(
        project_key=project_key, issuetype_name=issuetype_name, summary=summary,
        description=description, fields=fields, assignee=assignee, priority=priority,
        components=components, attachments=attachments, links=links, project_id=project_id)


def update_issue(account, password, issue_key, description=None, summary=None,
                 assignee=None, fields=None, project_key="IKQDM551EA", issuetype_name="ST-BUG"):
    """
    更新已有 JIRA 问题的快捷入口（封装登录 + 编辑）。详见 JIRAClient.update_issue。
    :return: True 表示有字段被更新
    """
    client = JIRAClient(account, password)
    if not client.login():
        raise Exception("JIRA登录失败")
    return client.update_issue(
        issue_key, description=description, summary=summary, assignee=assignee,
        fields=fields, project_key=project_key, issuetype_name=issuetype_name)


# 示例使用
if __name__ == "__main__":
    # 示例配置
    ACCOUNT = "venus.li@ikotek.com"
    PASSWORD = "@@@@@Aa13106680957"
    JQL = "issuetype = ST-BUG AND text ~ 'QDM565' ORDER BY status ASC"
    VERSION = "LTE01R07A01_C_SDK_E_QDM565_01.001.01.003_V04"

    # ---- 推荐：先拉全部JIRA列表，再逐个取完整详情，落盘为完整数据集 ----
    # 默认会自动追加 status != "ST_Closed"，因此这里拿到的就是你网页上看到的非关闭问题。
    try:
        res = export_complete_dataset(
            ACCOUNT, PASSWORD, JQL,
            output_dir=".", version=VERSION,
            concurrency=4, output_format="both", include_full_details=False)
        print(f"\n完整数据集 JSON: {res['json']}")
        print(f"完整数据集报告: {res['markdown']} / {res['docx']}")
        print(f"统计: {res['stats']}")
    except Exception as e:
        print(f"\n❌ 完整数据集导出失败: {e}")
        import traceback
        traceback.print_exc()

    # ---- 显式要包含 ST_Closed 时（如要进行全量历史统计）----
    # try:
    #     res_all = export_complete_dataset(
    #         ACCOUNT, PASSWORD, JQL,
    #         output_dir=".", version=VERSION,
    #         exclude_closed=False,          # 不追加 status != "ST_Closed"
    #         output_format="markdown")
    #     print(f"含关闭问题的报告: {res_all['markdown']}")
    # except Exception as e:
    #     print(f"导出失败: {e}")

    # ---- 基于已落盘的完整数据集，做任意后续操作（无需再访问 JIRA）----
    try:
        dataset = JIRACompleteDataset.load_json("jira_complete_dataset.json")
        only_major = [d for d in dataset
                      if JIRACompleteDataset._match_label(
                          d.get('custom_fields', {}),
                          JIRACompleteDataset.SEVERITY_LABEL_CANDIDATES) == 'B-Major']
        print(f"\n从完整数据集筛选 B-Major 共 {len(only_major)} 个")
        md2 = generate_complete_markdown(only_major, version="仅B-Major",
                                         output_file="JIRA_B-Major子集.md")
        print(f"子集报告: {md2}")
    except Exception as e:
        print(f"\n❌ 基于数据集的后续操作失败: {e}")
        import traceback
        traceback.print_exc()